from re import template
from docx import Document
import streamlit as st
import os
import io

st.image('zizzl health logo 22.png')

st.title("Proposal Tool")

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)

Client = st.text_input('Client/Prospect name')    

Broker = st.text_input('Brokerage Name')

Agent = st.text_input('Broker/Consultant Name')

Date = st.text_input('Date of Proposal')

Deductible = st.text_input('CSA Benchmark Plan - Deductible')

Out_of_Pocket = st.text_input('CSA Benchmark Plan - MAX Out of Pocket')

Carrier = st.text_input('CSA Benchmark Plan - Carrier')

Example = st.text_input('CSA Benchmark Plan - Premium for a non-smoking 30- year old in the baseline county')

Website = st.text_input('Broker Website')

Email = st.text_input('Broker Email')

Phone = st.text_input('Broker Phone')

template_file_path = 'Proposal - Copy.docx'

guardian = st.radio('Will this group utilize Guardian products?', ['Yes', 'No'])
premier = st.radio('Premier?', ['Yes', 'No'])

if st.button('SUBMIT'):
    
    variables = {
        "CLIENT NAME": Client,
        "CARRIER NAME": Carrier,
        "INSERT DATE HERE": Date, 
        "BROKER NAME": Broker,
        "AGENT NAME": Agent,
        "WEBSITE": Website,
        "EMAIL": Email,
        "PHONE": Phone,
        "DEDUCTIBLE": Deductible,
        "OOP": Out_of_Pocket,
        "EXAMPLE": Example,
        "PREMIER": 'Concierge Team assistance for technical support', 
        "GUARDIAN": 'Guardian Life Insurance Company product administration, including APIs connecting the systems with ease'
    }

    if (guardian == 'No'):

        variables['GUARDIAN'] = ''
    
    if (premier == 'No'):

        variables['PREMIER'] = ''



    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    docx_stream = io.BytesIO()
    template_document.save(docx_stream)
    docx_bytes = docx_stream.getvalue()

    st.download_button(
        label = "Download Proposal",
        data = docx_bytes,
        file_name = 'Proposal.docx',
        mime='application/msword'
    )










